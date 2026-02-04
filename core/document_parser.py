import pdfplumber
import docx
import io
import re

class DocumentParser:
    @staticmethod
    def parse_file(uploaded_file):
        """
        Universal Parser: Handles PDF and DOCX with Hindi/English support.
        Returns: (text, error_message)
        """
        try:
            file_type = uploaded_file.name.split('.')[-1].lower()
            text = ""

            # --- PDF HANDLING ---
            if file_type == 'pdf':
                with pdfplumber.open(uploaded_file) as pdf:
                    for page in pdf.pages:
                        # extract_text() handles Hindi unicode better than older libraries
                        page_text = page.extract_text() or ""
                        text += page_text + "\n"
                
                # Fallback: If pdfplumber got nothing (e.g., scanned image), warn the user
                if len(text.strip()) < 50:
                    return None, "⚠️ This PDF appears to be a scanned image. Please upload a digital PDF with selectable text."

            # --- DOCX HANDLING ---
            elif file_type in ['docx', 'doc']:
                doc = docx.Document(uploaded_file)
                # Reads paragraphs AND tables (important for legal docs)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"

            else:
                return None, "❌ Unsupported file format. Please upload PDF or DOCX."

            # --- CLEANING STAGE ---
            # Remove excessive whitespace but keep Hindi characters intact
            text = re.sub(r'\n+', '\n', text)  # Collapse multiple newlines
            text = re.sub(r'\s+', ' ', text)   # Collapse multiple spaces
            
            return text.strip(), None

        except Exception as e:
            return None, f"❌ Error parsing document: {str(e)}"